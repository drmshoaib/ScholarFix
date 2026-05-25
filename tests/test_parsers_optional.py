"""Optional parser smoke tests."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest


def test_tex_parser_extracts_plain_text(tex_upload: BytesIO) -> None:
    pytest.importorskip("pylatexenc")
    from parsers.tex_parser import extract_tex_text

    text = extract_tex_text(tex_upload)

    assert "Abstract" in text
    assert "smith2024" in text
    assert "notation score" in text


def test_docx_parser_extracts_paragraph_text() -> None:
    pytest.importorskip("docx")
    from docx import Document
    from parsers.docx_parser import extract_docx_text

    buffer = BytesIO()
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(buffer)
    buffer.seek(0)

    text = extract_docx_text(buffer)

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_pdf_report_generation_smoke() -> None:
    pytest.importorskip("reportlab")
    from reports.pdf_generator import generate_custom_pdf

    pdf_path = generate_custom_pdf(["Math issue"], ["Citation issue"])

    assert pdf_path.endswith(".pdf")
    assert Path(pdf_path).exists()
    Path(pdf_path).unlink(missing_ok=True)
